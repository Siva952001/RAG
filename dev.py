import streamlit as st
#import time
from googletrans import Translator
from langchain.chains import RetrievalQA
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.callbacks.manager import CallbackManager
from langchain.llms import Ollama
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain_community.vectorstores import Chroma
from langchain.embeddings.ollama import OllamaEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import pandas as pd
import os
#import shutil
import json
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from database import get_mongo_client, get_database, get_collection, save_chat_to_db


# Custom CSS for styling
st.markdown("""
    <style>
    .sidebar .sidebar-content {
        background-color: #f0f2f6;
    }
    .sidebar .sidebar-content h1 {
        font-size: 24px;
        color: #1f77b4;
    }
    .stTextInput>div>div>input {
        border-radius: 4px;
        border: 1px solid #ddd;
        padding: 10px;
    }
    .stButton>button {
        background-color: #1f77b4;
        color: white;
        border: none;
        padding: 10px 20px;
        border-radius: 4px;
    }
    .stButton>button:hover {
        background-color: #155a8a;
    }
    .chat-container {
        border: 1px solid #ddd;
        border-radius: 4px;
        padding: 10px;
        background-color: #fff;
    }
    .message {
        margin-bottom: 10px;
    }
    .message.user {
        color: #1f77b4;
    }
    .message.assistant {
        color: #e75f5f;
    }
    .message.translated {
        font-style: italic;
        color: #6c757d;
    }
    </style>
""", unsafe_allow_html=True)


# MongoDB setup
client = get_mongo_client()
db = get_database(client)
collection = get_collection(db)


MODEL_DIR = "C:/Users/DEV-037/.ollama/models/manifests/registry.ollama.ai/library"
FILE_DIR = 'files'
DATA_DIR = 'data'
SESSION_DIR = 'sessions'
translator = Translator()


os.makedirs(FILE_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(SESSION_DIR, exist_ok=True)


def list_models(model_dir):
    return [model for model in os.listdir(model_dir) if os.path.isdir(os.path.join(model_dir, model))]


def save_session(session_id):
    session_data = {
        "chat_history": st.session_state.chat_history,
        "processed_files": st.session_state.processed_files,
        "selected_model": st.session_state.selected_model
        
    }
    with open(os.path.join(SESSION_DIR, f"{session_id}.json"), 'w') as f:
        json.dump(session_data, f)
        
        
def load_session(session_id):
    with open(os.path.join(SESSION_DIR, f"{session_id}.json"), 'r') as f:
        session_data = json.load(f)
    st.session_state.chat_history = session_data.get("chat_history", [])
    st.session_state.processed_files = session_data.get("processed_files", {})
    st.session_state.selected_model = session_data.get("selected_model", None)
    
    
# def clear_data():
#     if os.path.exists(FILE_DIR):
#         shutil.rmtree(FILE_DIR)
#     # if os.path.exists(DATA_DIR):
#     #     shutil.rmtree(DATA_DIR)
#     os.makedirs(FILE_DIR, exist_ok=True)
#     os.makedirs(DATA_DIR, exist_ok=True)
#     st.session_state.chat_history = []
#     st.session_state.processed_files = {}


available_models = list_models(MODEL_DIR)


if 'template' not in st.session_state:
    st.session_state.template = """You are a knowledgeable chatbot, here to help with questions of the user. Your tone should be professional and informative.

    Context: {context}
    History: {history}

    User: {question}
    Chatbot:"""
if 'prompt' not in st.session_state:
    st.session_state.prompt = PromptTemplate(
        input_variables=["history", "context", "question"],
        template=st.session_state.template,
        
    )
if 'memory' not in st.session_state:
    st.session_state.memory = ConversationBufferMemory(
        memory_key="history",
        return_messages=True,
        input_key="question")
    

if 'processed_files' not in st.session_state:
    st.session_state.processed_files = {}
    

# Session management UI
st.sidebar.header("Chat Session Management")
session_id = st.sidebar.text_input("Session ID", value="default")
if st.sidebar.button("Save Session"):
    save_session(session_id)
    st.sidebar.success("Session saved successfully")
if st.sidebar.button("Load Session"):
    load_session(session_id)
    st.sidebar.success("Session loaded successfully")
# if st.sidebar.button("Clear"):
#     clear_data()
    st.sidebar.success("Cleared files and data successfully")
    
    
selected_model = st.selectbox("Select LLM Model", available_models)


if 'llm' not in st.session_state or st.session_state.selected_model != selected_model:
    st.session_state.selected_model = selected_model
    st.session_state.llm = Ollama(base_url="http://localhost:11434",
                                  model=selected_model,
                                  verbose=True,
                                  callback_manager=CallbackManager(
                                      [StreamingStdOutCallbackHandler()]),
                                  )
    st.session_state.embedding_model = OllamaEmbeddings(model=selected_model)
    

if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
    

st.sidebar.header("Upload Files")
uploaded_files = st.sidebar.file_uploader("Upload your files", type=['xlsx', 'xls', 'pdf', 'docx'], accept_multiple_files=True)


def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, "rb") as file:
        reader = PdfReader(file)
        for page_num in range(len(reader.pages)):
            text += reader.pages[page_num].extract_text() + "\n"
    return text


def extract_text_from_docx(file_path):
    text = ""
    doc = DocxDocument(file_path)
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


if uploaded_files is not None:
    for uploaded_file in uploaded_files:
        file_key = f"{uploaded_file.name}_{uploaded_file.size}"
        file_path = os.path.join(FILE_DIR, uploaded_file.name)
        vectorstore_path = os.path.join(DATA_DIR, f"{uploaded_file.name}_vectorstore")

        if file_key in st.session_state.processed_files:
            st.sidebar.warning(f"File '{uploaded_file.name}' already exists. You can ask questions.")
        elif os.path.exists(vectorstore_path):
            st.sidebar.warning(f"It looks like the vector store for '{uploaded_file.name}' is already present in our records. You can continue your chat.")
            st.session_state.vectorstore = Chroma(persist_directory=vectorstore_path, embedding_function=st.session_state.embedding_model)
            st.session_state.retriever = st.session_state.vectorstore.as_retriever()
            st.session_state.processed_files[file_key] = True
        else:
            try:
                with st.spinner("Analyzing your document..."):
                    # Save the uploaded file to the file directory
                    with open(file_path, "wb") as f:
                        f.write(uploaded_file.getvalue())

                    # Extract text based on file type
                    if uploaded_file.name.endswith('.xlsx') or uploaded_file.name.endswith('.xls'):
                        xls = pd.ExcelFile(file_path)
                        all_content = ""
                        for sheet_name in xls.sheet_names:
                            df = pd.read_excel(xls, sheet_name=sheet_name)
                            data = df.to_dict(orient='records')
                            content = "\n".join([str(row) for row in data])
                            all_content += content + "\n"
                    elif uploaded_file.name.endswith('.pdf'):
                        all_content = extract_text_from_pdf(file_path)
                    elif uploaded_file.name.endswith('.docx'):
                        all_content = extract_text_from_docx(file_path)
                    else:
                        st.sidebar.warning(f"Unsupported file type: {uploaded_file.name}")
                        continue
                    
                    # Process and store the text
                    text_splitter = RecursiveCharacterTextSplitter(
                        chunk_size=2000,
                        chunk_overlap=300,
                        length_function=len
                    )
                    splits = text_splitter.split_text(all_content)
                    documents = [Document(page_content=split) for split in splits]

                    st.session_state.vectorstore = Chroma.from_documents(
                        documents=documents,
                        embedding=st.session_state.embedding_model,
                        persist_directory=vectorstore_path
                    )
                    st.session_state.vectorstore.persist()

                    st.session_state.processed_files[file_key] = True
                    st.sidebar.success(f"File '{uploaded_file.name}' processed successfully!")

                st.session_state.retriever = st.session_state.vectorstore.as_retriever()
                

            except PermissionError:
                st.error(f"Permission denied: Unable to write to {file_path}. Please check if the file is open or locked.")
                
                
    if 'vectorstore' in st.session_state:
        st.session_state.retriever = st.session_state.vectorstore.as_retriever()
        

        if 'qa_chain' not in st.session_state:
            st.session_state.qa_chain = RetrievalQA.from_chain_type(
                llm=st.session_state.llm,
                chain_type='stuff',
                retriever=st.session_state.retriever,
                verbose=True,
                chain_type_kwargs={
                    "verbose": True,
                    "prompt": st.session_state.prompt,
                    "memory": st.session_state.memory,
                }
            )

st.write("---")

# Layout containers
chat_container = st.container()
input_container = st.container()

# Display Chat History
with chat_container:
    st.subheader("Chat")
    for message in st.session_state.chat_history:
        if message['role'] == 'user':
            st.markdown(f"<div class='chat-container'><p class='message user'>{message['message']}</p></div>", unsafe_allow_html=True)
        else:
            st.markdown(f"<div class='chat-container'><p class='message assistant'>{message['message']}</p></div>", unsafe_allow_html=True)
            st.markdown(f"<div class='chat-container'><p class='message translated'>{message['translated_message']}</p></div>", unsafe_allow_html=True)
            

# Input Field at the Bottom
with input_container:
    st.write("---")
    user_input = st.text_input(
        "Ask a question:",
        key="user_input",
        placeholder="Type your question here..."
    )
    

def safe_translate(text, dest_lang):
    try:
        translated = translator.translate(text, dest=dest_lang)
        return translated.text
    except Exception as e:
        st.error(f"Translation failed: {str(e)}")
        return text
    

def save_to_mongo(user_input, response_text, translated_response):
    from datetime import datetime
    current_time = datetime.now().strftime('%d-%m-%Y')
    chat_data = {
        "user_query": user_input,
        "assistant_response": response_text,
        "translated_response": translated_response,
        "timestamp": current_time
    }
    save_chat_to_db(collection, chat_data)
    

if st.button("Submit"):
    if user_input:
        user_message = {"role": "user", "message": user_input}
        st.session_state.chat_history.append(user_message)
        st.markdown(f"<div class='chat-container'><p class='message user'>{user_input}</p></div>", unsafe_allow_html=True)

        with st.spinner("Assistant is typing..."):
            # Pass the input as a dictionary with 'query' key
            response = st.session_state.qa_chain({"query": user_input})
            response_text = response['result']

            translated_response = safe_translate(response_text, 'ta')
            
            # MongoDB
            save_to_mongo(user_input, response_text, translated_response)

            assistant_message = {"role": "assistant", "message": response_text, "translated_message": translated_response}
            st.session_state.chat_history.append(assistant_message)

            st.markdown(f"<div class='chat-container'><p class='message assistant'>{response_text}</p></div>", unsafe_allow_html=True)
            st.markdown(f"<div class='chat-container'><p class='message translated'>{translated_response}</p></div>", unsafe_allow_html=True)
            

# Date input field in the sidebar
date_input = st.sidebar.date_input("Enter Date to view previous chats", min_value=None)


def retrieve_data_by_date(date):
    # Format the date as '%d-%m-%Y' for MongoDB query
    formatted_date = date.strftime('%d-%m-%Y')
    query = {"timestamp": formatted_date}
    return list(collection.find(query))


# Add button to load data by date
if st.sidebar.button("Submit" , key = "submit_button"):
    if date_input:
        # Retrieve and display data
        formatted_date = date_input.strftime('%d-%m-%Y')
        data = retrieve_data_by_date(date_input)
        
        if data:
            st.sidebar.subheader(f"Data for {formatted_date}")
            for record in data:
                st.sidebar.write(f"**User Query:** {record['user_query']}")
                st.sidebar.write(f"**Assistant Response:** {record['assistant_response']}")
                st.sidebar.write(f"**Translated Response:** {record['translated_response']}")
                st.sidebar.write(f"**Timestamp:** {record['timestamp']}")
                st.sidebar.write("---")
        else:
            st.sidebar.write(f"No records found for {formatted_date}.")
    else:
        st.sidebar.write("Please select a date.")
        


