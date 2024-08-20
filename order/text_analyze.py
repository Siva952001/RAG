import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.vectorstores import Chroma
from langchain.embeddings.ollama import OllamaEmbeddings
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain.schema import Document
from googletrans import Translator
import os
from langchain.llms import Ollama
from langchain.text_splitter import RecursiveCharacterTextSplitter

MODEL_DIR = "C:/Users/DEV-037/.ollama/models/manifests/registry.ollama.ai/library"
FILE_DIR = 'files'
DATA_DIR = 'data'

def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, "rb") as file:
        reader = PdfReader(file)
        for page_num in range(len(reader.pages)):
            text += reader.pages[page_num].extract_text() + "\n"
    return text

def extract_text_from_docx(file_path):
    text = ""
    doc = DocxDocument(file_path)
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def process_file(uploaded_file, file_path, model_name):
    # Save the uploaded file to the file directory
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getvalue())

    # Extract text based on file type
    all_content = ""
    if uploaded_file.name.endswith('.xlsx') or uploaded_file.name.endswith('.xls'):
        xls = pd.ExcelFile(file_path)
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            data = df.to_dict(orient='records')
            content = "\n".join([str(row) for row in data])
            all_content += content + "\n"
    elif uploaded_file.name.endswith('.pdf'):
        all_content = extract_text_from_pdf(file_path)
    elif uploaded_file.name.endswith('.docx'):
        all_content = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type")

    # Process and store the text
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=300,
        length_function=len
    )
    splits = text_splitter.split_text(all_content)
    documents = [Document(page_content=split) for split in splits]

    vectorstore_path = os.path.join(DATA_DIR, f"{uploaded_file.name}_vectorstore")
    vectorstore = Chroma.from_documents(
        documents=documents,
        embedding=OllamaEmbeddings(model=model_name),
        persist_directory=vectorstore_path
    )
    vectorstore.persist()

def translate_text(text, dest_lang):
    translator = Translator()
    try:
        translated = translator.translate(text, dest=dest_lang)
        return translated.text
    except Exception as e:
        return text  # Fallback to the original text if translation fails

def get_vectorstore(model_name):
    vectorstore_path = os.path.join(DATA_DIR, f"{model_name}_vectorstore")
    vectorstore = Chroma(persist_directory=vectorstore_path, embedding_function=OllamaEmbeddings(model=model_name))
    retriever = vectorstore.as_retriever()
    prompt = PromptTemplate(
        input_variables=["history", "context", "question"],
        template="""You are a knowledgeable chatbot, here to help with questions of the user. Your tone should be professional and informative.
        Context: {context}
        History: {history}
        User: {question}
        Chatbot:"""
    )
    memory = ConversationBufferMemory(
        memory_key="history",
        return_messages=True,
        input_key="question"
    )
    qa_chain = RetrievalQA.from_chain_type(
        llm=Ollama(base_url="http://localhost:11434", model=model_name),
        chain_type='stuff',
        retriever=retriever,
        verbose=True,
        chain_type_kwargs={
            "verbose": True,
            "prompt": prompt,
            "memory": memory,
        }
    )
    return vectorstore, qa_chain
